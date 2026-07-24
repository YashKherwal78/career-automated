import os
import re
import docx
from pypdf import PdfReader
from typing import Dict, Any

class DocumentTextExtractor:
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text content from .pdf, .docx, or .txt files."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return DocumentTextExtractor._extract_from_pdf(file_path)
        elif ext == ".docx":
            return DocumentTextExtractor._extract_from_docx(file_path)
        elif ext in (".txt", ".md"):
            return DocumentTextExtractor._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        text_content = []
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
            return "\n".join(text_content).strip()
        except Exception as e:
            raise RuntimeError(f"Error reading PDF file: {str(e)}")

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_content.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text]
                    if row_text:
                        text_content.append(" | ".join(row_text))
            return "\n".join(text_content).strip()
        except Exception as e:
            raise RuntimeError(f"Error reading DOCX file: {str(e)}")

    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception as e:
            raise RuntimeError(f"Error reading TXT file: {str(e)}")
