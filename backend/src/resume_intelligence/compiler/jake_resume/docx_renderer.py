"""
Jake Resume DOCX Renderer Subsystem.

Generates editable DOCX documents preserving Jake Resume typography and geometry:
- 0.5in margins
- Bold header with centered contact bar
- Uppercase section titles with bottom borders
- Bold titles and italic company/dates
- Bullet list formatting
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from src.resume_intelligence.compiler.jake_resume.models import StructuredResume


class JakeDOCXRenderer:
    """DOCX Renderer for Jake Resume V1."""

    def render(self, resume: StructuredResume, docx_path: str) -> str:
        out_dir = os.path.dirname(docx_path)
        os.makedirs(out_dir, exist_ok=True)

        doc = Document()

        # Page Setup — 0.5 inch margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Style Helpers
        def add_name_header(name: str):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(name.upper())
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = 'Arial'
            p.paragraph_format.space_after = Pt(2)

        def add_contact_line(contact_text: str):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(contact_text)
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            p.paragraph_format.space_after = Pt(10)

        def add_section_title(title: str):
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            
            # Bottom border emulation xml
            pBrd = parse_xml(f'<w:pBrd {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBrd>')
            p._p.get_or_add_pPr().append(pBrd)

        def add_two_column_line(left_text: str, right_text: str, is_bold: bool = True):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            
            r_left = p.add_run(left_text)
            r_left.bold = is_bold
            r_left.font.size = Pt(9.5)
            r_left.font.name = 'Arial'

            # Tab stop right aligned at 7.5 inches
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), 2)
            r_left.add_text('\t')
            
            r_right = p.add_run(right_text)
            r_right.font.size = Pt(9)
            r_right.font.name = 'Arial'

        def add_bullet(bullet_text: str):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(bullet_text)
            run.font.size = Pt(9)
            run.font.name = 'Arial'

        # 1. Render Header
        add_name_header(resume.name)
        c_text = f"{resume.contact.phone}  |  {resume.contact.email}  |  {resume.contact.linkedin}  |  {resume.contact.github}"
        add_contact_line(c_text)

        # 2. Dynamic Section Ordering
        for sec in resume.section_order:
            if sec == "summary" and resume.summary:
                add_section_title("Professional Summary")
                p = doc.add_paragraph(resume.summary)
                p.paragraph_format.space_after = Pt(4)

            elif sec == "education" and resume.education:
                add_section_title("Education")
                for edu in resume.education:
                    add_two_column_line(edu.institution, f"{edu.start_date} – {edu.end_date}", is_bold=True)
                    add_two_column_line(f"{edu.degree} in {edu.field_of_study}", edu.location, is_bold=False)

            elif sec == "experience" and resume.experience:
                add_section_title("Experience")
                for exp in resume.experience:
                    add_two_column_line(exp.title, f"{exp.start_date} – {exp.end_date}", is_bold=True)
                    add_two_column_line(exp.company, exp.location, is_bold=False)
                    for b in exp.bullets:
                        add_bullet(b)

            elif sec == "projects" and resume.projects:
                add_section_title("Projects")
                for proj in resume.projects:
                    tech_str = f" | {', '.join(proj.technologies)}" if proj.technologies else ""
                    add_two_column_line(f"{proj.title}{tech_str}", proj.date, is_bold=True)
                    for b in proj.bullets:
                        add_bullet(b)

            elif sec == "skills" and resume.skill_categories:
                add_section_title("Technical Skills")
                for cat in resume.skill_categories:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    r_cat = p.add_run(f"{cat.category_name}: ")
                    r_cat.bold = True
                    r_cat.font.size = Pt(9)
                    r_cat.font.name = 'Arial'

                    r_skills = p.add_run(", ".join(cat.skills))
                    r_skills.font.size = Pt(9)
                    r_skills.font.name = 'Arial'

        doc.save(docx_path)
        return docx_path
