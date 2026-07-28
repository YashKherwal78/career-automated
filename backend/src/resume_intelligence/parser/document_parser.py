"""
Decoupled Multi-Format Resume Parser & Extraction Subsystem (Module 2).

Parser decouples Raw Document Extraction from Normalization.
Supported Formats:
- Text PDF (pypdf stream extractor)
- Scanned / Image PDF (via OCR Engine fallback)
- DOCX (native python-docx element parser)
- TXT / Markdown (structural line block parser)
"""

import os
import re
from typing import Dict, Any, List, Optional
from pydantic import BaseModel, Field
from pypdf import PdfReader
try:
    import docx
except ImportError:
    docx = None


class RawBlock(BaseModel):
    block_type: str = "text"  # 'header', 'paragraph', 'bullet', 'contact', 'table'
    content: str
    page_number: int = 1
    bounding_box: Optional[List[float]] = None


class RawExtraction(BaseModel):
    file_path: str
    file_type: str  # 'pdf_text', 'pdf_scanned', 'docx', 'txt', 'md'
    is_scanned: bool = False
    raw_text: str = ""
    blocks: List[RawBlock] = Field(default_factory=list)
    extraction_confidence: float = 1.0


class DocumentParser:
    """Multi-format Document Parser returning RawExtraction objects."""

    def parse_document(self, file_path: str) -> RawExtraction:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume document not found at {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext == ".docx":
            return self._parse_docx(file_path)
        elif ext in [".txt", ".md"]:
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _parse_pdf(self, file_path: str) -> RawExtraction:
        try:
            reader = PdfReader(file_path)
            full_text = []
            blocks = []
            page_num = 1

            for page in reader.pages:
                text = page.extract_text() or ""
                full_text.append(text)
                lines = [l.strip() for l in text.split("\n") if l.strip()]
                for line in lines:
                    b_type = "bullet" if line.startswith(("-", "•", "*", "–")) else "text"
                    if len(line) < 40 and line.isupper():
                        b_type = "header"
                    blocks.append(RawBlock(block_type=b_type, content=line, page_number=page_num))
                page_num += 1

            combined_text = "\n".join(full_text)
            
            # Detect if PDF is scanned (very low character count per page)
            is_scanned = len(combined_text.strip()) < 100

            return RawExtraction(
                file_path=file_path,
                file_type="pdf_scanned" if is_scanned else "pdf_text",
                is_scanned=is_scanned,
                raw_text=combined_text,
                blocks=blocks,
                extraction_confidence=0.4 if is_scanned else 0.98
            )
        except Exception as e:
            return RawExtraction(
                file_path=file_path,
                file_type="pdf_text",
                is_scanned=True,
                raw_text="",
                blocks=[],
                extraction_confidence=0.1
            )

    def _parse_docx(self, file_path: str) -> RawExtraction:
        if not docx:
            raise ImportError("python-docx package is required for DOCX parsing")
        
        doc = docx.Document(file_path)
        full_text = []
        blocks = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            full_text.append(text)
            b_type = "bullet" if p.style.name.startswith("List") or text.startswith(("-", "•")) else "text"
            if p.style.name.startswith("Heading"):
                b_type = "header"
            blocks.append(RawBlock(block_type=b_type, content=text, page_number=1))

        # Tables parsing
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    full_text.append(row_text)
                    blocks.append(RawBlock(block_type="table", content=row_text, page_number=1))

        combined = "\n".join(full_text)

        return RawExtraction(
            file_path=file_path,
            file_type="docx",
            is_scanned=False,
            raw_text=combined,
            blocks=blocks,
            extraction_confidence=0.99
        )

    def _parse_txt(self, file_path: str) -> RawExtraction:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        lines = [l.strip() for l in content.split("\n") if l.strip()]
        blocks = []

        for line in lines:
            b_type = "header" if line.startswith("#") or (len(line) < 30 and line.isupper()) else "text"
            if line.startswith(("-", "*", "•")):
                b_type = "bullet"
            blocks.append(RawBlock(block_type=b_type, content=line, page_number=1))

        return RawExtraction(
            file_path=file_path,
            file_type="txt" if file_path.endswith(".txt") else "md",
            is_scanned=False,
            raw_text=content,
            blocks=blocks,
            extraction_confidence=1.0
        )
